'''from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from pydantic import BaseModel
import google.generativeai as genai
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import hashlib
from docx import Document
from PIL import Image
import io

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
else:
    raise ValueError("Google API Key not found. Please check your environment variables.")

# Initialize FastAPI app
app = FastAPI()

# Pydantic model for job description input
class JobDescription(BaseModel):
    job_description: str

# Helper function for caching extracted text per file
def hash_file(file_content: bytes) -> str:
    md5_hash = hashlib.md5(file_content).hexdigest()
    return md5_hash

# Extract text from PDF
def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_lines = [page.extract_text() for page in pdf_reader.pages]
        return "".join(text_lines)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {e}")

# Extract text from DOCX
def extract_text_from_docx(file_content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {e}")

# Extract text from image (OCR)
def extract_text_from_image(file_content: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(file_content))
        # Use OCR library like pytesseract here (not implemented in this example)
        return "Text extracted from image (OCR not implemented in this example)."
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading image: {e}")

# Process uploaded files and extract text
def process_files(files: list[UploadFile]) -> str:
    combined_text = ""
    for file in files:
        file_content = file.file.read()
        if file.filename.endswith(".pdf"):
            combined_text += extract_text_from_pdf(file_content) + "\n\n"
        elif file.filename.endswith(".docx"):
            combined_text += extract_text_from_docx(file_content) + "\n\n"
        elif file.filename.endswith((".png", ".jpg", ".jpeg")):
            combined_text += extract_text_from_image(file_content) + "\n\n"
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.filename}")
    return combined_text

# Generate analysis using Gemini model
def generate_analysis(text: str, job_description: str) -> str:
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
    prompt = (
        "Assess candidate fit for the job description. Consider substitutes for skills, experience, match percentage in tabular form:\n\n"
        "Skills: Match or equivalent technologies.\n"
        "Experience: Relevance to key responsibilities.\n"
        "Fit: Suitability based on experience and skills.\n\n"
        f"Job Description:\n{job_description}\n\nResume Content:\n{text}"
    )
    response = model.generate_content([prompt], stream=True)
    response.resolve()
    return response.text

# API endpoint for file upload and analysis
@app.post("/analyze")
async def analyze(
    job_description: str = Form(...),  # Accept job_description as form data
    files: list[UploadFile] = File(...)
):
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded.")
    
    # Process files and extract text
    combined_text = process_files(files)
    
    # Generate analysis
    analysis_result = generate_analysis(combined_text, job_description)
    
    # Return JSON response
    return {
        "status": "success",
        "analysis": analysis_result
    }

# Run the FastAPI app
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)'''

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from pydantic import BaseModel
import google.generativeai as genai
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import hashlib
from docx import Document
from PIL import Image
import io
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

if not GOOGLE_API_KEY:
    logger.error("Google API Key not found in environment variables")
    raise ValueError("Google API Key not found. Please check your .env file.")

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    logger.info("Successfully configured Gemini AI")
except Exception as e:
    logger.error(f"Failed to configure Gemini AI: {str(e)}")
    raise

# Initialize FastAPI app
app = FastAPI(title="Resume Analyzer API",
             description="API for analyzing resumes against job descriptions using Gemini AI")

# Constants
ALLOWED_FILE_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'image/png': 'png',
    'image/jpeg': 'jpg'
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

class JobDescription(BaseModel):
    job_description: str

def hash_file(file_content: bytes) -> str:
    try:
        md5_hash = hashlib.md5(file_content).hexdigest()
        return md5_hash
    except Exception as e:
        logger.error(f"Hashing failed: {str(e)}")
        raise HTTPException(status_code=500, detail="File processing error")

def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_lines = []
        for page in pdf_reader.pages:
            try:
                text = page.extract_text()
                if text:
                    text_lines.append(text)
            except Exception as page_error:
                logger.warning(f"Error extracting text from PDF page: {str(page_error)}")
                continue
        return "\n".join(text_lines) if text_lines else ""
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def extract_text_from_image(file_content: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(file_content))
        # Placeholder - in production, use pytesseract or similar
        return "Image text extraction would appear here (OCR not implemented in this example)"
    except Exception as e:
        logger.error(f"Image processing failed: {str(e)}")
        raise HTTPException(status_code=400, detail=f"Error reading image: {str(e)}")

def process_files(files: list[UploadFile]) -> str:
    combined_text = []
    for file in files:
        try:
            logger.info(f"Processing file: {file.filename}")
            
            # Validate file size
            file.file.seek(0, 2)  # Seek to end
            file_size = file.file.tell()
            file.file.seek(0)  # Reset pointer
            if file_size > MAX_FILE_SIZE:
                raise HTTPException(status_code=413, detail=f"File too large: {file.filename}")
            
            file_content = file.file.read()
            
            # Validate file type
            if file.content_type not in ALLOWED_FILE_TYPES:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file.content_type}"
                )
            
            # Process based on type
            if file.content_type == 'application/pdf':
                text = extract_text_from_pdf(file_content)
            elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                text = extract_text_from_docx(file_content)
            else:  # Image
                text = extract_text_from_image(file_content)
            
            if text.strip():
                combined_text.append(text)
                logger.info(f"Successfully processed {file.filename}")
            else:
                logger.warning(f"No text extracted from {file.filename}")
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing {file.filename}: {str(e)}"
            )
    
    return "\n\n".join(combined_text) if combined_text else ""

def generate_analysis(text: str, job_description: str) -> str:
    try:
        logger.info("Generating analysis with Gemini")
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        
        prompt = f"""
        Analyze the candidate's resume against the job description and provide:
        1. Skills match (including comparable technologies)
        2. Experience relevance
        3. Overall fit percentage
        
        Format the response in clear sections with bullet points.
        Include a table for skills matching.
        
        Job Description:
        {job_description}
        
        Resume Content:
        {text}
        """
        
        response = model.generate_content(prompt)
        return response.text
    
    except Exception as e:
        logger.error(f"Analysis generation failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Analysis generation failed: {str(e)}"
        )

@app.post("/analyze")
async def analyze(
    job_description: str = Form(..., description="The job description to analyze against"),
    files: list[UploadFile] = File(..., description="Resume files (PDF, DOCX, or images)")
):
    """
    Analyze resumes against a job description using Gemini AI.
    
    Returns:
    - Skills match analysis
    - Experience relevance
    - Overall fit assessment
    """
    try:
        if not job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        if not files:
            raise HTTPException(status_code=400, detail="No files uploaded")
        
        logger.info(f"Starting analysis for job description: {job_description[:50]}...")
        combined_text = process_files(files)
        
        if not combined_text.strip():
            raise HTTPException(
                status_code=400,
                detail="No readable text could be extracted from the provided files"
            )
        
        analysis_result = generate_analysis(combined_text, job_description)
        
        return {
            "status": "success",
            "analysis": analysis_result,
            "text_length": len(combined_text)
        }
        
    except HTTPException as he:
        logger.error(f"HTTP error in analysis: {he.detail}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error in analysis: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred: {str(e)}"
        )

@app.get("/health")
async def health_check():
    """Endpoint for health checks"""
    return {"status": "healthy", "api": "resume-analyzer"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")